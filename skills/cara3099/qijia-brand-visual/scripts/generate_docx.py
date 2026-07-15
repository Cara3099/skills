#!/usr/bin/env python3
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from validate_input import load_and_validate


NAVY = "172A4A"
RED = "CE2A2A"
INK = "15181D"
MUTED = "68717F"
LINE = "D8DCE2"


def rgb(value):
    return RGBColor.from_string(value)


def style_run(run, size=11, color=INK, bold=False, font="Microsoft YaHei"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold


def set_paragraph_border(paragraph, side="bottom", color=NAVY, size=10, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def add_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def next_id(nodes, attribute):
    values = []
    for node in nodes:
        value = node.get(qn(attribute))
        if value is not None and str(value).isdigit():
            values.append(int(value))
    return max(values, default=0) + 1


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = next_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId")
    num_id = next_id(numbering.findall(qn("w:num")), "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    level.extend([start, num_fmt, level_text, level_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def add_numbered_list(doc, items, bold_prefix=None, size=11, line_spacing=1.5, space_after=5):
    num_id = create_numbering(doc)
    for item in items:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, num_id)
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.35)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.line_spacing = line_spacing
        if bold_prefix and isinstance(item, tuple):
            style_run(paragraph.add_run(item[0]), size, INK, True)
            style_run(paragraph.add_run(item[1]), size, INK)
        else:
            style_run(paragraph.add_run(str(item)), size, INK)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    for name, size, color, before, after in (
        ("Title", 26, NAVY, 0, 12),
        ("Heading 1", 16, NAVY, 18, 9),
        ("Heading 2", 13.5, NAVY, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def populate_header(header, logo_path, brand, date_value):
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Cm(17))
    table.autofit = False
    table.columns[0].width = Cm(12)
    table.columns[1].width = Cm(5)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left.add_run().add_picture(str(logo_path), width=Cm(0.55))
    style_run(left.add_run(f"  {brand}"), 9.5, NAVY, True)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(right.add_run(date_value), 8.5, MUTED)
    set_paragraph_border(left, "bottom", RED, 7, 4)
    set_paragraph_border(right, "bottom", RED, 7, 4)


def populate_footer(footer):
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(paragraph.add_run("广东企加税务师事务所  |  "), 8, MUTED)
    add_field(paragraph, "PAGE")


def add_header_footer(doc, section, logo_path, brand, date_value):
    doc.settings.odd_and_even_pages_header_footer = True
    populate_header(section.header, logo_path, brand, date_value)
    populate_header(section.even_page_header, logo_path, brand, date_value)
    populate_footer(section.footer)
    populate_footer(section.even_page_footer)


def add_cover(doc, data):
    meta = data["meta"]
    top = doc.add_paragraph()
    top.paragraph_format.space_before = Pt(34)
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(top.add_run("财税政策文件"), 12, RED, True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(22)
    style_run(title.add_run(meta["title"]), 28, NAVY, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = Cm(0)
    subtitle.paragraph_format.space_before = Pt(12)
    subtitle.paragraph_format.space_after = Pt(48)
    style_run(subtitle.add_run(meta["subtitle"]), 14, MUTED)

    line = doc.add_paragraph()
    line.paragraph_format.first_line_indent = Cm(0)
    set_paragraph_border(line, "bottom", NAVY, 12, 8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.first_line_indent = Cm(0)
    info.paragraph_format.space_before = Pt(28)
    style_run(info.add_run(f'{meta["issue"]}\n\n编制日期：{meta["date"]}\n\n适用对象：{meta["audience"]}'), 11, INK)

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company.paragraph_format.first_line_indent = Cm(0)
    company.paragraph_format.space_before = Pt(52)
    style_run(company.add_run(meta["brand"]), 13, NAVY, True)
    doc.add_page_break()


def add_body_intro(doc, data):
    meta = data["meta"]
    heading = doc.add_paragraph("一、编制说明", style="Heading 1")
    set_paragraph_border(heading, "bottom", NAVY, 8, 4)
    paragraph = doc.add_paragraph(
        f'本简报根据截至{meta["date"]}公开发布的财税政策和征管提示编制，'
        "用于帮助企业管理层、财务及税务负责人了解近期事项、判断可能影响并安排内部落实工作。"
        "文中政策摘要以公开文件为基础，执行时仍应结合具体交易事实、企业实际情况和主管机关现行口径。"
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading = doc.add_paragraph("二、近期重点事项", style="Heading 1")
    set_paragraph_border(heading, "bottom", NAVY, 8, 4)
    items = []
    for item in data["summary"]:
        items.append((f'{item["label"]}（{item["value"]}）：', item["detail"]))
    add_numbered_list(doc, items, bold_prefix=True)


def add_policy(doc, policy, index):
    heading = doc.add_paragraph(f'{index}. {policy["title"]}', style="Heading 2")
    set_paragraph_border(heading, "bottom", LINE, 5, 3)

    metadata = doc.add_paragraph()
    metadata.paragraph_format.first_line_indent = Cm(0)
    style_run(metadata.add_run("文件信息："), 10, NAVY, True)
    style_run(
        metadata.add_run(
            f'{policy["document_no"]}；发布日期 {policy["published"]}；实施或关注日期 {policy["effective"]}；'
            f'责任角色 {policy["owner"]}；建议时限 {policy["deadline"]}。'
        ),
        10,
        MUTED,
    )

    overview = doc.add_paragraph()
    style_run(overview.add_run("政策概述："), 11, INK, True)
    style_run(overview.add_run(policy["summary"]), 11, INK)

    sub = doc.add_paragraph("（一）对企业的主要影响", style="Heading 3")
    add_numbered_list(doc, policy["impacts"])
    sub = doc.add_paragraph("（二）建议落实事项", style="Heading 3")
    add_numbered_list(doc, policy["actions"])

    source = doc.add_paragraph()
    source.paragraph_format.first_line_indent = Cm(0)
    source.paragraph_format.space_after = Pt(10)
    style_run(source.add_run("官方来源："), 9, NAVY, True)
    style_run(source.add_run(f'{policy["source_title"]}\n{policy["source_url"]}'), 8.5, MUTED)


def add_actions_and_sources(doc, data):
    heading = doc.add_paragraph("四、内部执行建议", style="Heading 1")
    set_paragraph_border(heading, "bottom", NAVY, 8, 4)
    items = []
    for action in data["actions"]:
        prefix = f'{action["action"]}（责任角色：{action["owner"]}；建议时限：{action["deadline"]}；当前状态：{action["status"]}）'
        items.append(prefix)
    add_numbered_list(doc, items)

    heading = doc.add_paragraph("五、资料来源及使用说明", style="Heading 1")
    set_paragraph_border(heading, "bottom", NAVY, 8, 4)
    source_items = [f'{p["source_title"]}，{p["document_no"]}，{p["source_url"]}' for p in data["policies"]]
    add_numbered_list(doc, source_items, size=9, line_spacing=1.18, space_after=2)

    disclaimer = doc.add_paragraph()
    disclaimer.paragraph_format.first_line_indent = Cm(0.74)
    disclaimer.paragraph_format.space_before = Pt(8)
    disclaimer.paragraph_format.space_after = Pt(8)
    style_run(disclaimer.add_run("使用说明："), 10.5, NAVY, True)
    style_run(disclaimer.add_run(data["meta"]["disclaimer"]), 10.5, INK)
    set_paragraph_border(disclaimer, "top", RED, 7, 5)


def main():
    data_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    skill_root = Path(sys.argv[3]).resolve() if len(sys.argv) > 3 else Path(__file__).resolve().parents[1]
    data = load_and_validate(data_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.65)
    section.right_margin = Cm(2.45)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = False

    configure_styles(doc)
    add_header_footer(doc, section, skill_root / "assets/brand/qijia-logo.png", data["meta"]["brand"], data["meta"]["date"])
    add_cover(doc, data)
    add_body_intro(doc, data)

    heading = doc.add_paragraph("三、政策内容", style="Heading 1")
    set_paragraph_border(heading, "bottom", NAVY, 8, 4)
    for index, policy in enumerate(data["policies"], 1):
        add_policy(doc, policy, index)
    add_actions_and_sources(doc, data)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = data["meta"]["title"]
    doc.core_properties.subject = data["meta"]["subtitle"]
    doc.core_properties.author = data["meta"]["brand"]
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
