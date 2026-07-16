#!/usr/bin/env python3
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from validate_input import load_and_validate


ORANGE = "FF5001"
INK = "20242C"
NAVY = "162033"
MUTED = "666D78"


def set_font(run, size=11, color=INK, bold=False, font="SimSun"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, end])


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in (
        ("Title", 20, 0, 12),
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 12, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc, data, logo_path):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path.exists():
        header.add_run().add_picture(str(logo_path), width=Cm(2.3))
    set_font(header.add_run(f"  {data['meta']['document_type']}"), 8.5, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run(f"{data.get('footer_note', '代理记账认准代理猫')}  ·  "), 8, MUTED)
    add_page_field(footer)


def add_title_block(doc, data):
    meta = data["meta"]
    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.first_line_indent = Cm(0)
    set_font(eyebrow.add_run(meta["document_type"]), 9, ORANGE, True, "Microsoft YaHei")

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.first_line_indent = Cm(0)
    set_font(title.add_run(meta["title"]), 20, NAVY, True, "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.first_line_indent = Cm(0)
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run(meta["subtitle"]), 11, MUTED, False, "Microsoft YaHei")

    metadata = doc.add_paragraph()
    metadata.paragraph_format.first_line_indent = Cm(0)
    metadata.paragraph_format.space_after = Pt(16)
    set_font(metadata.add_run(f"日期：{meta['date']}    对象：{meta['audience']}"), 9, MUTED)


def add_body(doc, data):
    doc.add_paragraph("概览", style="Heading 1")
    doc.add_paragraph(data.get("overview", ""))

    for item in data.get("highlights", []):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.keep_together = True
        set_font(p.add_run(f"{item.get('title', '重点')}："), 11, INK, True)
        set_font(p.add_run(item.get("detail", "")), 11, INK)

    for section in data.get("sections", []):
        doc.add_paragraph(section.get("heading", "未命名章节"), style="Heading 1")
        for text in section.get("paragraphs", []):
            doc.add_paragraph(text)
        for text in section.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            set_font(p.add_run(text), 11, INK)

    if data.get("actions"):
        doc.add_paragraph("后续行动", style="Heading 1")
        for item in data["actions"]:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.keep_together = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(f"{item.get('id', '')}  {item.get('title', '')}"), 11, NAVY, True, "Microsoft YaHei")
            detail = item.get("detail", "")
            owner = item.get("owner", "")
            if detail:
                set_font(p.add_run(f"\n{detail}"), 11, INK)
            if owner:
                set_font(p.add_run(f"\n责任：{owner}"), 9, MUTED)


def main():
    if len(sys.argv) < 3:
        raise SystemExit("用法：generate_docx.py <content.json> <output.docx> [skill-root]")
    data_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    skill_root = Path(sys.argv[3]).resolve() if len(sys.argv) > 3 else Path(__file__).resolve().parents[1]
    data = load_and_validate(data_path)
    doc = Document()
    configure_document(doc)
    add_header_footer(doc, data, skill_root / "assets/brand/agentcat-logo.png")
    add_title_block(doc, data)
    add_body(doc, data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
